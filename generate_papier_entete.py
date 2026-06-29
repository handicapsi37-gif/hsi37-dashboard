#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Générateur du papier à en-tête HSI37
Produit : docs/modeles/papier-entete-HSI37.docx
Usage   : python3 generate_papier_entete.py
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ──────────────────────────────────────────────────────────────────
#  Charte graphique HSI37
# ──────────────────────────────────────────────────────────────────
BLEU      = RGBColor(0x3B, 0x77, 0xB4)
TEXTE     = RGBColor(0x40, 0x3E, 0x3E)
JAUNE_HEX = 'F7CD46'

DIR    = os.path.dirname(os.path.abspath(__file__))


def detecter_police():
    """Retourne 'Montserrat' si installée sur le système, sinon 'Helvetica'."""
    dirs = [
        '/Library/Fonts',
        os.path.expanduser('~/Library/Fonts'),
        '/System/Library/Fonts',
    ]
    for d in dirs:
        try:
            if any('Montserrat' in f for f in os.listdir(d)):
                return 'Montserrat'
        except OSError:
            pass
    return 'Helvetica'

POLICE_CORPS = detecter_police()
LOGO   = os.path.join(DIR, 'assets', 'hsi37-redim-demi.png')
OUTPUT = os.path.join(DIR, 'docs', 'modeles', 'papier-entete-HSI37.docx')


# ──────────────────────────────────────────────────────────────────
#  Utilitaires XML
# ──────────────────────────────────────────────────────────────────
def emu_twips(emu):
    return str(int(emu * 1440 // 914400))


def vider_container(container):
    """Supprime tous les paragraphes et tableaux du container (en-tête/pied)."""
    el = container._element
    for child in list(el):
        tag = child.tag.split('}')[-1]
        if tag in ('p', 'tbl'):
            el.remove(child)


def supprimer_bordures_tableau(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    borders = OxmlElement('w:tblBorders')
    for nom in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{nom}')
        b.set(qn('w:val'), 'nil')
        borders.append(b)
    tblPr.append(borders)


def largeur_cellule(cell, emu):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), emu_twips(emu))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)
    borders = OxmlElement('w:tcBorders')
    for nom in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{nom}')
        b.set(qn('w:val'), 'nil')
        borders.append(b)
    tcPr.append(borders)


def fond_cellule(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def espacement(para, avant=0, apres=0):
    pPr = para._p.get_or_add_pPr()
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), str(avant))
    sp.set(qn('w:after'), str(apres))
    pPr.append(sp)


def run_txt(para, texte, taille=10, couleur=None, gras=False, police=None):
    run = para.add_run(texte)
    run.font.size = Pt(taille)
    run.bold = gras
    run.font.name = police or POLICE_CORPS
    if couleur:
        run.font.color.rgb = couleur
    return run


# ──────────────────────────────────────────────────────────────────
#  Filet jaune — bordure de paragraphe (plus fiable que tableau dans Word)
# ──────────────────────────────────────────────────────────────────
def bordure_jaune(para, position='bottom', epaisseur_pt=2):
    """Ajoute une bordure jaune en haut ou en bas d'un paragraphe."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    el = OxmlElement(f'w:{position}')
    el.set(qn('w:val'), 'single')
    el.set(qn('w:sz'), str(int(epaisseur_pt * 8)))  # unités 1/8 pt
    el.set(qn('w:space'), '0')
    el.set(qn('w:color'), JAUNE_HEX)
    pBdr.append(el)
    pPr.append(pBdr)


# ──────────────────────────────────────────────────────────────────
#  En-tête
# ──────────────────────────────────────────────────────────────────
def construire_entete(section):
    header = section.header
    header.is_linked_to_previous = False
    vider_container(header)

    # Tableau 2 colonnes : logo (gauche) | coordonnées (droite)
    table = header.add_table(rows=1, cols=2, width=Cm(16))
    supprimer_bordures_tableau(table)

    W_LOGO  = Cm(5)
    W_TEXTE = Cm(11)
    cell_logo  = table.rows[0].cells[0]
    cell_texte = table.rows[0].cells[1]
    largeur_cellule(cell_logo, W_LOGO)
    largeur_cellule(cell_texte, W_TEXTE)

    # Colonne gauche — logo
    p_logo = cell_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    espacement(p_logo, 0, 0)
    if os.path.exists(LOGO):
        p_logo.add_run().add_picture(LOGO, width=Cm(4.2))

    # Colonne droite — nom association
    p_nom = cell_texte.paragraphs[0]
    p_nom.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    espacement(p_nom, 0, 40)
    run_txt(p_nom, "Handicap Solidarité pour l'Inclusion 37",
            taille=11, couleur=BLEU, gras=True)

    for texte, taille in [
        ("17 rue Gabriel Péri, 37700 Saint-Pierre-des-Corps", 9),
        ("07 43 29 58 30  ·  handicapsi37@gmail.com  ·  www.hsi37.fr", 9),
        ("Association loi 1901  ·  N° RNA : W372020254  ·  SIRET : 93164414000010", 8),
    ]:
        p = cell_texte.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        espacement(p, 0, 20)
        run_txt(p, texte, taille=taille, couleur=TEXTE)

    # Filet jaune sous l'en-tête — bordure basse 2pt
    p_filet = header.add_paragraph()
    espacement(p_filet, 80, 0)
    bordure_jaune(p_filet, position='bottom', epaisseur_pt=2)


# ──────────────────────────────────────────────────────────────────
#  Pied de page
# ──────────────────────────────────────────────────────────────────
def construire_pied(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    vider_container(footer)

    # Filet jaune au-dessus du pied — bordure haute 2pt
    p_filet = footer.add_paragraph()
    espacement(p_filet, 0, 40)
    bordure_jaune(p_filet, position='top', epaisseur_pt=2)

    # Coordonnées complètes centrées
    p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    espacement(p, 80, 20)
    run_txt(
        p,
        "Association Handicap Solidarité pour l'Inclusion 37  ·  "
        "17 rue Gabriel Péri, 37700 Saint-Pierre-des-Corps  ·  "
        "07 43 29 58 30  ·  handicapsi37@gmail.com  ·  www.hsi37.fr",
        taille=8, couleur=TEXTE,
    )

    # Numéro de page centré
    p_page = footer.add_paragraph()
    p_page.alignment = WD_ALIGN_PARAGRAPH.CENTER
    espacement(p_page, 0, 0)
    run = p_page.add_run()
    run.font.size = Pt(8)
    run.font.color.rgb = TEXTE
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.extend([fldChar, instr, fldChar2, fldChar3])


# ──────────────────────────────────────────────────────────────────
#  Main
# ──────────────────────────────────────────────────────────────────
def main():
    doc = Document()
    section = doc.sections[0]

    section.page_width      = Cm(21)
    section.page_height     = Cm(29.7)
    section.left_margin     = Cm(2.5)
    section.right_margin    = Cm(2.5)
    section.top_margin      = Cm(4.0)
    section.bottom_margin   = Cm(2.5)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)

    construire_entete(section)
    construire_pied(section)

    # Corps : paragraphe vierge prêt à rédiger
    para = doc.add_paragraph()
    espacement(para, 0, 200)
    run = para.add_run()
    run.font.name = POLICE_CORPS
    run.font.size = Pt(11)
    run.font.color.rgb = TEXTE

    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
    doc.save(OUTPUT)
    print(f"✓ Papier à en-tête généré : {OUTPUT}")


if __name__ == '__main__':
    main()
